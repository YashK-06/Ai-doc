from docx import Document


def extract_document(file_path):
    doc = Document(file_path)

    document_data = {
        "paragraphs": [],
        "tables": []
    }

    # Extract paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            document_data["paragraphs"].append(text)

    # Extract tables
    for table in doc.tables:
        table_data = []

        for row in table.rows:
            row_data = []

            for cell in row.cells:
                row_data.append(cell.text.strip())

            table_data.append(row_data)

        document_data["tables"].append(table_data)

    return document_data